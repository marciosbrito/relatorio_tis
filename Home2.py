from pathlib import Path
from docx2txt import process
from docx.shared import Inches
import pandas as pd
import relatorios.security.segredos as ss
from PIL import Image
import random
from PIL import Image
import docx
import streamlit as st
from docx import Document
import io
from os import path
import os


st.sidebar.header("Faça se Login abaixo")
# Lista de usuários e senhas
user_pwd = ss.usuarios_e_senhas

# Função para autenticar o usuário


def autenticar_usuario(usuario, senha):
    senha_correta = user_pwd.get(usuario)
    if senha_correta == senha:
        return True
    return False


# Criar uma interface de login com Streamlit
# Campos de entrada para nome de usuário e senha
usuario = st.sidebar.text_input("Usuário")
senha = st.sidebar.text_input("Senha", type="password")

# Botão de login
if st.sidebar.checkbox("Entrar"):
    if autenticar_usuario(usuario, senha):
        st.sidebar.success("Login bem-sucedido!")
        st.subheader('Bem vindo ao Relatório de Atendimento Técnico')
        # Coleta dados do usuário
        aleatorio = random.randint(1000, 99999)
        num_rat = st.subheader(f'RAT Nº: {aleatorio}')
        motivo = st.text_input('Motivo do Chamado')
        cliente = st.text_input('Cliente')
        local = st.text_input("Local:")
        tipo = st.selectbox("Selecione o Tipo de Manutenção", [
                            "Preventiva", "Corretiva", 'Vistoria'])
        tecnico1 = st.text_input('Nome do Técnico 1')
        tecnico2 = st.text_input('Nome do Técnico 2')
        data = st.date_input('Digite a Data', format=('DD/MM/YYYY'))
        h_inicio = st.time_input('Hora Inicial')
        h_final = st.time_input('Hora Final')
        cabecalho_descricao = st.text('Descrição das Atividades:')
        descricao = st.text_area('')
        nome_arq = st.text_input(
            "Digite o nome do relatoriocom (.doc no final)")
        foto1 = st.file_uploader("Faça o upload das imagens1 (JPEG ou PNG)", type=[
            "jpg", "jpeg", "png"])
        texto1 = st.text_input("Digite descrição da Imagem1")
        foto2 = st.file_uploader("Faça o upload das imagens2 (JPEG ou PNG)", type=[
            "jpg", "jpeg", "png"])
        texto2 = st.text_input("Digite descrição da Imagem2")
        foto3 = st.file_uploader("Faça o upload das imagens3 (JPEG ou PNG)", type=[
            "jpg", "jpeg", "png"])
        texto3 = st.text_input("Digite descrição da Imagem3")
        if st.button("Pré Visualização do Relatório"):
            st.write(f':blue[RAT Nº:  ] {aleatorio}')
            st.write(f':blue[Motivo:  ]   {motivo}')
            st.write(f':blue[Cliente:  ]   {cliente}')
            st.write(f':blue[Local:  ]   {local}')
            st.write(f':blue[Tipo:  ]   {tipo}')
            st.write(f':blue[Técnico 1:  ]   {tecnico1}')
            st.write(f':blue[Técnico 2:  ]   {tecnico2}')
            st.write(f':blue[Data:  ]   {data}')
            st.write(f':blue[Hora Inicial:  ]   {h_inicio}')
            st.write(f':blue[Hora Final:  ]   {h_final}')
            st.write(f':blue[Descrição das Atividades:  ]   {descricao}')
            # st.image(f'Imagem 1: {foto1}')
            # st.image(f'Imagem 2: {foto2}')
            # st.image(f'Imagem 3: {foto3}')

        if st.button("Gerar Relatório"):
            # Cria um documento Word
            doc = Document()

        # Logo do Relatorio
            doc.add_picture(
                'C:\\Users\\Administrador\\Documents\\app\\app_relatorio_tis\\relatorios\\Imagens\\Logo_Telefonica2.png')
            doc.add_heading('Relatório de Atendimento Técnico', 0)

        # Adiciona os dados do formulário ao documento
            doc.add_paragraph(f"Número do RAT:  {aleatorio}")
            doc.add_paragraph(
                f"Motivo do Chamado: {motivo.upper()}", style='Intense Quote')
            doc.add_paragraph(
                f"Cliente: {cliente.upper()}",  style='Body Text')
            doc.add_paragraph(f"Local: {local.upper()}",  style='Body Text')
            doc.add_paragraph(
                f"Motivo do Chamado: {tipo.upper()}",  style='Body Text')
            doc.add_paragraph(
                f"Técnico 1: {tecnico1.upper()}",  style='Body Text')
            doc.add_paragraph(
                f"Técnico 2: {tecnico2.upper()}",  style='Body Text')
            doc.add_paragraph(
                f"Data do Atendimento: {data}",  style='Body Text')
            doc.add_paragraph(
                f"inicio da Atividade: {h_inicio}",  style='Body Text')
            doc.add_paragraph(
                f"Termino da Atividade: {h_final}",  style='Body Text')
            doc.add_paragraph('Descrição das Atividades:', style='Body Text')
            doc.add_paragraph(f"{descricao.upper()}",  style='Heading 1')
            doc.add_paragraph(
                '------------------------------------------------------------------------------------------')
            if texto1:
                doc.add_paragraph(f'Imagem 1\n{texto1}')
            if foto1:
                ler_imagem1 = foto1.read()
                doc.add_picture(io.BytesIO(ler_imagem1),
                                width=Inches(4), height=Inches(3))

            if texto2:
                doc.add_paragraph(f'Imagem 2\n{texto2}')
            if foto2:
                ler_imagem2 = foto2.read()
                doc.add_picture(io.BytesIO(ler_imagem2),
                                width=Inches(4), height=Inches(3))

            if texto3:
                doc.add_paragraph(f'Imagem 3\n{texto3}')
            if foto3:
                ler_imagem3 = foto3.read()
                doc.add_picture(io.BytesIO(ler_imagem3),
                                width=Inches(4), height=Inches(3))

        # Salva o documento Word
            doc_file = (nome_arq)
            doc.save(
                f'C:\\Users\\Administrador\\Documents\\app\\app_relatorio_tis\\relatorios\\relatorios_tis{doc_file}')

            st.success(
                f"Documento Word gerado com sucesso: [{doc_file}]({doc_file})")
    else:
        st.error('Credenciais incorretas. Tente novamente')
